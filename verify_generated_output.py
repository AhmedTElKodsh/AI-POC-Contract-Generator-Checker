"""
Verify Generated Output Script
Checks that generated DOCX files contain correct extracted values
"""

import sys
import os
from docx import Document

def verify_docx_content(file_path):
    """Verify that DOCX contains expected values from sample RFP"""
    print(f"\n{'='*60}")
    print(f"Verifying: {file_path}")
    print('='*60)
    
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return False
    
    try:
        doc = Document(file_path)
        
        # Extract all text from document
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Expected values from sample_client_rfp.docx
        expected_values = {
            "Project Name": "New Capital Water Treatment Plant Extension",
            "Client": "New Urban Communities Authority (NUCA)",
            "Location": "New Administrative Capital, Egypt",
            "Duration": "18"  # months
        }
        
        print("\n📋 Checking for expected values:")
        all_found = True
        
        for field, expected in expected_values.items():
            if expected in full_text:
                print(f"✅ {field}: '{expected}' found")
            else:
                print(f"❌ {field}: '{expected}' NOT FOUND")
                all_found = False
        
        # Check for glossary terms
        print("\n📚 Checking for glossary terms:")
        glossary_terms = ["drainage", "flood", "water"]
        found_terms = []
        for term in glossary_terms:
            if term.lower() in full_text.lower():
                found_terms.append(term)
                print(f"✅ Term '{term}' found")
        
        if found_terms:
            print(f"✅ Found {len(found_terms)} glossary terms")
        else:
            print("⚠️  No glossary terms found (may be expected if KB didn't match)")
        
        # Check for references
        print("\n📖 Checking for technical references:")
        if "ECP" in full_text or "Egyptian Code" in full_text:
            print("✅ Technical references found")
        else:
            print("⚠️  No technical references found")
        
        # File statistics
        print(f"\n📊 Document Statistics:")
        print(f"   Paragraphs: {len(doc.paragraphs)}")
        print(f"   Total text length: {len(full_text)} characters")
        print(f"   File size: {os.path.getsize(file_path)} bytes")
        
        if all_found:
            print(f"\n✅ VERIFICATION PASSED: All expected values found!")
            return True
        else:
            print(f"\n⚠️  VERIFICATION INCOMPLETE: Some values missing")
            return False
        
    except Exception as e:
        print(f"❌ Error reading document: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("\n" + "="*60)
    print("🔍 GENERATED OUTPUT VERIFICATION")
    print("="*60)
    
    # Files to verify
    files_to_check = [
        "proposals/test_output.docx",
        "proposals/e2e_test_output.docx"
    ]
    
    results = {}
    for file_path in files_to_check:
        results[file_path] = verify_docx_content(file_path)
    
    # Summary
    print("\n" + "="*60)
    print("📊 VERIFICATION SUMMARY")
    print("="*60)
    
    passed = sum(1 for v in results.values() if v)
    total = len(results)
    
    for file_path, result in results.items():
        status = "✅ PASS" if result else "⚠️  INCOMPLETE"
        print(f"{status} - {os.path.basename(file_path)}")
    
    print(f"\n🎯 Results: {passed}/{total} files verified")
    
    if passed == total:
        print("\n🎉 All generated files contain correct extracted values!")
        return 0
    else:
        print(f"\n⚠️  {total - passed} file(s) need review")
        return 1

if __name__ == "__main__":
    exit_code = main()
    sys.exit(exit_code)
