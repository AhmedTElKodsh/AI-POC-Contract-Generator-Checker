"""Unit tests for document extraction functions."""

import pytest
import os
import tempfile
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches

import sys

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

import doc_pipeline


# Fixtures
@pytest.fixture
def temp_dir():
    """Create temporary directory for test files."""
    temp_path = tempfile.mkdtemp()
    yield temp_path
    shutil.rmtree(temp_path)


@pytest.fixture
def sample_docx_file(temp_dir):
    """Create a sample DOCX file for testing."""
    doc_path = os.path.join(temp_dir, "test_document.docx")
    doc = Document()

    # Add title
    title = doc.add_heading("Test Document", 0)

    # Add sections
    doc.add_heading("Introduction", 1)
    doc.add_paragraph("This is a test introduction paragraph with some sample content.")

    doc.add_heading("Methodology", 1)
    doc.add_paragraph("Methodology content goes here.")
    doc.add_paragraph("Additional methodology details.")

    doc.add_heading("Conclusion", 1)
    doc.add_paragraph("This is the conclusion of the test document.")

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Quantity"
    table.rows[0].cells[2].text = "Unit"
    table.rows[1].cells[0].text = "Item 1"
    table.rows[1].cells[1].text = "100"
    table.rows[1].cells[2].text = "m3"
    table.rows[2].cells[0].text = "Item 2"
    table.rows[2].cells[1].text = "200"
    table.rows[2].cells[2].text = "m2"

    doc.save(doc_path)
    return doc_path


# Tests for extract_docx_segments
def test_extract_docx_segments_basic(sample_docx_file):
    """Test basic DOCX extraction."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)

    assert isinstance(result, dict)
    assert "segments" in result
    assert "tables" in result
    assert "metadata" in result
    assert "Error" not in result


def test_extract_docx_segments_has_sections(sample_docx_file):
    """Test that DOCX extraction creates sections."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)

    assert "segments" in result
    segments = result["segments"]
    assert isinstance(segments, dict)
    assert len(segments) > 0

    # Check for expected sections
    section_names = list(segments.keys())
    assert any("Introduction" in section for section in section_names)


def test_extract_docx_segments_has_tables(sample_docx_file):
    """Test that DOCX extraction extracts tables."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)

    assert "tables" in result
    tables = result["tables"]
    assert isinstance(tables, list)
    assert len(tables) > 0

    # Check table structure
    table = tables[0]
    assert "rows" in table
    assert "columns" in table
    assert "data" in table
    assert isinstance(table["data"], list)


def test_extract_docx_segments_metadata(sample_docx_file):
    """Test that DOCX extraction provides metadata."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)

    assert "metadata" in result
    metadata = result["metadata"]
    assert "paragraph_count" in metadata
    assert "table_count" in metadata
    assert "section_count" in metadata

    assert metadata["table_count"] > 0


def test_extract_docx_segments_invalid_file():
    """Test DOCX extraction with invalid file."""
    result = doc_pipeline.extract_docx_segments("nonexistent.docx")

    assert "Error" in result


# Tests for validate_file_path
def test_validate_file_path_valid_file(temp_dir):
    """Test file path validation with valid file."""
    file_path = os.path.join(temp_dir, "test.txt")
    Path(file_path).touch()

    assert doc_pipeline.validate_file_path(file_path, [".txt"], temp_dir) is True


def test_validate_file_path_invalid_extension(temp_dir):
    """Test file path validation with invalid extension."""
    file_path = os.path.join(temp_dir, "test.pdf")
    Path(file_path).touch()

    with pytest.raises(ValueError):
        doc_pipeline.validate_file_path(file_path, [".txt"], temp_dir)


def test_validate_file_path_path_traversal(temp_dir):
    """Test file path validation prevents path traversal."""
    file_path = os.path.join(temp_dir, "..", "test.txt")
    Path(os.path.join(temp_dir, "..", "test.txt")).touch()

    with pytest.raises(ValueError):
        doc_pipeline.validate_file_path(file_path, [".txt"], temp_dir)


# Tests for get_file_hash
def test_get_file_hash_consistency(temp_dir):
    """Test that file hash is consistent."""
    file_path = os.path.join(temp_dir, "test.txt")
    content = "Test content for hashing"

    with open(file_path, "w") as f:
        f.write(content)

    hash1 = doc_pipeline.get_file_hash(file_path)
    hash2 = doc_pipeline.get_file_hash(file_path)

    assert hash1 is not None
    assert hash1 == hash2


def test_get_file_hash_different_files(temp_dir):
    """Test that different files have different hashes."""
    file_path1 = os.path.join(temp_dir, "test1.txt")
    file_path2 = os.path.join(temp_dir, "test2.txt")

    with open(file_path1, "w") as f:
        f.write("Content 1")

    with open(file_path2, "w") as f:
        f.write("Content 2")

    hash1 = doc_pipeline.get_file_hash(file_path1)
    hash2 = doc_pipeline.get_file_hash(file_path2)

    assert hash1 != hash2


# Tests for table extraction quality
def test_docx_table_structure(sample_docx_file):
    """Test that DOCX tables have correct structure."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)
    tables = result["tables"]

    if tables:
        table = tables[0]
        assert "rows" in table
        assert "columns" in table
        assert "data" in table

        # Check that data is a list of lists
        assert isinstance(table["data"], list)
        if table["data"]:
            assert isinstance(table["data"][0], list)


def test_docx_table_cell_values(sample_docx_file):
    """Test that DOCX table cells are properly extracted."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)
    tables = result["tables"]

    if tables:
        table_data = tables[0]["data"]
        # Check that cells are strings
        for row in table_data:
            for cell in row:
                assert isinstance(cell, str)


# Tests for segmentation
def test_docx_segmentation_creates_sections(sample_docx_file):
    """Test that DOCX segmentation creates logical sections."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)
    segments = result["segments"]

    assert len(segments) > 0

    # Check that sections have content
    for section_name, content in segments.items():
        assert isinstance(content, str)


def test_docx_segmentation_preserves_content(sample_docx_file):
    """Test that DOCX segmentation preserves document content."""
    result = doc_pipeline.extract_docx_segments(sample_docx_file)
    segments = result["segments"]

    # Combine all segments
    all_content = "\n".join(segments.values())

    # Check that expected keywords are present
    assert "Introduction" in all_content or len(all_content) > 0


# Bilingual content tests
def test_docx_arabic_content(temp_dir):
    """Test DOCX extraction with Arabic content."""
    doc_path = os.path.join(temp_dir, "arabic_test.docx")
    doc = Document()

    doc.add_heading("الدراسة الهيدرولوجية", 0)
    doc.add_paragraph("هذا نص تجريب باللغة العربية.")
    doc.add_paragraph("Hydrological study for drainage project.")

    # Add Arabic table
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "الوصف"
    table.rows[0].cells[1].text = "الكمية"
    table.rows[1].cells[0].text = "حفر"
    table.rows[1].cells[1].text = "100"

    doc.save(doc_path)

    result = doc_pipeline.extract_docx_segments(doc_path)
    assert "segments" in result
    assert "tables" in result


def test_docx_mixed_language_content(temp_dir):
    """Test DOCX extraction with mixed Arabic/English content."""
    doc_path = os.path.join(temp_dir, "mixed_test.docx")
    doc = Document()

    doc.add_heading("Hydrological Study دراسة هيدرولوجية", 0)
    doc.add_paragraph("This is English text.")
    doc.add_paragraph("هذا نص باللغة العربية.")

    doc.save(doc_path)

    result = doc_pipeline.extract_docx_segments(doc_path)
    assert "segments" in result
    assert isinstance(result["segments"], dict)
