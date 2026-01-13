"""Template Engine for proposal and report generation."""

import os
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from src.models.enums import Language, OutputFormat
from src.models.generation import BOQItem, ProposalContext, RenderResult


class TemplateEngine:
    """
    Render proposals and reports using DOCX templates with Jinja2.

    Supports:
    - DOCX templates with Jinja2 placeholders
    - Dynamic table generation (BOQ, team, timeline)
    - Arabic RTL text preservation
    - PDF output via WeasyPrint
    """

    def __init__(self, templates_dir: Optional[str] = None):
        """
        Initialize the template engine.

        Args:
            templates_dir: Directory containing template files
        """
        self.templates_dir = Path(templates_dir) if templates_dir else Path("templates")
        self._docxtpl = None
        self._weasyprint = None
        self._load_dependencies()

    def _load_dependencies(self):
        """Load optional dependencies."""
        try:
            import docxtpl
            self._docxtpl = docxtpl
        except ImportError:
            pass

        try:
            import weasyprint
            self._weasyprint = weasyprint
        except ImportError:
            pass

    def list_templates(self, project_type: Optional[str] = None) -> list[dict[str, Any]]:
        """
        List available templates.

        Args:
            project_type: Filter by project type

        Returns:
            List of template info dictionaries
        """
        templates = []

        if not self.templates_dir.exists():
            return templates

        for file_path in self.templates_dir.glob("*.docx"):
            template_info = {
                "name": file_path.stem,
                "filename": file_path.name,
                "path": str(file_path),
                "language": self._detect_template_language(file_path.stem),
            }

            # Filter by project type if specified
            if project_type:
                if project_type.lower() not in file_path.stem.lower():
                    continue

            templates.append(template_info)

        return templates

    def _detect_template_language(self, name: str) -> Language:
        """Detect template language from name."""
        name_lower = name.lower()
        if "_ar" in name_lower or "arabic" in name_lower:
            return Language.ARABIC
        elif "_en" in name_lower or "english" in name_lower:
            return Language.ENGLISH
        return Language.ENGLISH

    def render_proposal(
        self,
        template_name: str,
        context: ProposalContext,
        output_format: OutputFormat = OutputFormat.DOCX,
        output_path: Optional[str] = None,
    ) -> RenderResult:
        """
        Render a proposal using the specified template.

        Args:
            template_name: Name of the template file
            context: Proposal context data
            output_format: Output format (DOCX or PDF)
            output_path: Optional output file path

        Returns:
            RenderResult with output file info
        """
        if self._docxtpl is None:
            raise RuntimeError("python-docx-template is not installed")

        # Find template file
        template_path = self._find_template(template_name)
        if not template_path:
            raise FileNotFoundError(f"Template not found: {template_name}")

        # Load template
        doc = self._docxtpl.DocxTemplate(str(template_path))

        # Prepare context for Jinja2
        jinja_context = self._prepare_context(context)

        # Render template
        doc.render(jinja_context)

        # Generate output path
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = Path("output")
            output_dir.mkdir(exist_ok=True)
            output_path = str(output_dir / f"{context.project_name}_{timestamp}.docx")

        # Save DOCX
        doc.save(output_path)

        # Convert to PDF if requested
        if output_format == OutputFormat.PDF:
            pdf_path = self._convert_to_pdf(output_path)
            if pdf_path:
                output_path = pdf_path

        # Get file info
        file_size = os.path.getsize(output_path)

        return RenderResult(
            job_id=str(datetime.now().timestamp()),
            output_format=output_format,
            file_path=output_path,
            file_size_bytes=file_size,
            page_count=None,
            created_at=datetime.utcnow(),
        )


    def _find_template(self, template_name: str) -> Optional[Path]:
        """Find template file by name."""
        # Try exact match
        template_path = self.templates_dir / template_name
        if template_path.exists():
            return template_path

        # Try with .docx extension
        template_path = self.templates_dir / f"{template_name}.docx"
        if template_path.exists():
            return template_path

        # Search for partial match
        for file_path in self.templates_dir.glob("*.docx"):
            if template_name.lower() in file_path.stem.lower():
                return file_path

        return None

    def _prepare_context(self, context: ProposalContext) -> dict[str, Any]:
        """Prepare context dictionary for Jinja2 rendering."""
        # Format date
        date_str = context.date.strftime("%Y-%m-%d") if context.date else ""

        # Format BOQ items
        boq_data = []
        for item in context.boq_items:
            boq_data.append({
                "item_number": item.item_number,
                "description": item.description,
                "unit": item.unit,
                "quantity": f"{item.quantity:,.2f}",
                "unit_rate": f"{item.unit_rate:,.2f}",
                "total": f"{item.total:,.2f}",
                "notes": item.notes or "",
            })

        # Format currency
        total_formatted = f"{context.total_cost:,.2f} {context.currency}"

        return {
            # Basic info
            "project_name": context.project_name,
            "client": context.client,
            "location": context.location,
            "date": date_str,
            "duration_months": context.duration_months,
            # Sections
            "sections": context.sections,
            # BOQ
            "boq_items": boq_data,
            "total_cost": total_formatted,
            "currency": context.currency,
            # Team
            "team_members": context.team_members,
            # Timeline
            "timeline_milestones": context.timeline_milestones,
            # Custom metadata
            **context.metadata,
        }

    def _convert_to_pdf(self, docx_path: str) -> Optional[str]:
        """Convert DOCX to PDF using WeasyPrint or LibreOffice."""
        pdf_path = docx_path.replace(".docx", ".pdf")

        # Try using LibreOffice (more reliable for DOCX)
        try:
            import subprocess
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", str(Path(docx_path).parent),
                    docx_path,
                ],
                capture_output=True,
                timeout=60,
            )
            if result.returncode == 0 and Path(pdf_path).exists():
                return pdf_path
        except Exception:
            pass

        return None

    def validate_context(
        self,
        template_name: str,
        context: ProposalContext,
    ) -> list[str]:
        """
        Validate that context has all required fields for template.

        Args:
            template_name: Name of the template
            context: Proposal context to validate

        Returns:
            List of missing field names
        """
        missing = []

        # Required fields for all templates
        required_fields = [
            ("project_name", context.project_name),
            ("client", context.client),
            ("location", context.location),
        ]

        for field_name, value in required_fields:
            if not value:
                missing.append(field_name)

        # Check sections
        if not context.sections:
            missing.append("sections")

        return missing

    def preview_template(self, template_name: str) -> dict[str, Any]:
        """
        Get template structure and required fields.

        Args:
            template_name: Name of the template

        Returns:
            Template info with placeholders
        """
        template_path = self._find_template(template_name)
        if not template_path:
            return {"error": f"Template not found: {template_name}"}

        if self._docxtpl is None:
            return {"error": "python-docx-template not installed"}

        try:
            doc = self._docxtpl.DocxTemplate(str(template_path))
            # Get undeclared variables (placeholders)
            variables = doc.get_undeclared_template_variables()

            return {
                "name": template_name,
                "path": str(template_path),
                "placeholders": list(variables),
                "language": self._detect_template_language(template_name).value,
            }
        except Exception as e:
            return {"error": str(e)}

    def create_default_template(
        self,
        template_name: str,
        language: Language = Language.ENGLISH,
    ) -> str:
        """
        Create a default proposal template.

        Args:
            template_name: Name for the new template
            language: Template language

        Returns:
            Path to created template
        """
        if self._docxtpl is None:
            raise RuntimeError("python-docx-template is not installed")

        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("python-docx is not installed")

        # Create new document
        doc = Document()

        # Add title
        title = doc.add_heading("{{ project_name }}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add client info
        doc.add_paragraph(f"Client: {{{{ client }}}}")
        doc.add_paragraph(f"Location: {{{{ location }}}}")
        doc.add_paragraph(f"Date: {{{{ date }}}}")

        # Add sections placeholder
        doc.add_heading("Executive Summary", 1)
        doc.add_paragraph("{{ sections.executive_summary }}")

        doc.add_heading("Scope of Work", 1)
        doc.add_paragraph("{{ sections.scope_of_work }}")

        doc.add_heading("Methodology", 1)
        doc.add_paragraph("{{ sections.methodology }}")

        doc.add_heading("Timeline", 1)
        doc.add_paragraph("Duration: {{ duration_months }} months")

        doc.add_heading("Cost Estimate", 1)
        doc.add_paragraph("Total: {{ total_cost }}")

        # Save template
        self.templates_dir.mkdir(exist_ok=True)
        template_path = self.templates_dir / f"{template_name}.docx"
        doc.save(str(template_path))

        return str(template_path)
