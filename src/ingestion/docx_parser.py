"""DOCX document parser using python-docx."""

from pathlib import Path
from typing import Optional

from src.ingestion.exceptions import CorruptedFileError, UnsupportedFormatError
from src.ingestion.language_detector import LanguageDetector
from src.models.documents import DocumentMetadata, DocumentSection, ParsedDocument, TableData
from src.models.enums import DocumentType, Language, SectionType


class DOCXParser:
    """Parse DOCX documents and extract structured content."""

    def __init__(self):
        """Initialize the DOCX parser."""
        self.language_detector = LanguageDetector()
        self._docx = None
        self._load_dependencies()

    def _load_dependencies(self):
        """Load optional dependencies."""
        try:
            import docx
            self._docx = docx
        except ImportError:
            pass

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a DOCX document and extract structured content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ParsedDocument with extracted content

        Raises:
            UnsupportedFormatError: If file is not a DOCX
            CorruptedFileError: If file cannot be parsed
        """
        path = Path(file_path)

        # Validate file exists and is DOCX
        if not path.exists():
            raise CorruptedFileError(file_path, "File does not exist")

        if path.suffix.lower() != ".docx":
            raise UnsupportedFormatError(file_path, path.suffix)

        if self._docx is None:
            raise CorruptedFileError(file_path, "python-docx library not available")

        try:
            doc = self._docx.Document(file_path)

            # Extract all text
            raw_text = self._extract_raw_text(doc)

            # Detect primary language
            primary_language = self.language_detector.detect(raw_text)

            # Extract sections from paragraphs
            sections = self._extract_sections(doc, primary_language)

            # Extract tables
            tables = self._extract_tables(doc)

            # Attach tables to relevant sections or create table sections
            self._attach_tables_to_sections(sections, tables)

            # Create metadata
            metadata = DocumentMetadata(
                document_id="",  # Will be set by caller
                filename=path.name,
                document_type=DocumentType.DOCX,
                language=primary_language,
                total_pages=None,  # DOCX doesn't have fixed pages
            )

            # Generate markdown representation
            markdown = self._generate_markdown(sections, tables)

            return ParsedDocument(
                metadata=metadata,
                sections=sections,
                raw_text=raw_text,
                markdown=markdown,
            )

        except Exception as e:
            if isinstance(e, (CorruptedFileError, UnsupportedFormatError)):
                raise
            raise CorruptedFileError(file_path, str(e))

    def _extract_raw_text(self, doc) -> str:
        """Extract all text from document."""
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)


    def _extract_sections(self, doc, primary_language: Language) -> list[DocumentSection]:
        """Extract sections from document paragraphs."""
        sections = []
        current_section = None
        current_content = []

        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style else ""

            # Check if this is a heading
            if style_name.startswith("Heading") or style_name.startswith("Title"):
                # Save previous section
                if current_section:
                    current_section.content = "\n".join(current_content).strip()
                    sections.append(current_section)
                    current_content = []

                title = paragraph.text.strip()
                section_type = self._detect_section_type(title)

                current_section = DocumentSection(
                    section_type=section_type,
                    title=title,
                    content="",
                    subsections=[],
                    tables=[],
                    page_start=None,
                    page_end=None,
                )
            else:
                # Regular paragraph
                if paragraph.text.strip():
                    current_content.append(paragraph.text)

        # Save last section
        if current_section:
            current_section.content = "\n".join(current_content).strip()
            sections.append(current_section)
        elif current_content:
            # No headings found, create single section
            sections.append(
                DocumentSection(
                    section_type=SectionType.OTHER,
                    title="Content",
                    content="\n".join(current_content).strip(),
                    subsections=[],
                    tables=[],
                    page_start=None,
                    page_end=None,
                )
            )

        return sections

    def _extract_tables(self, doc) -> list[TableData]:
        """Extract all tables from document."""
        tables = []

        for table in doc.tables:
            headers = []
            rows = []

            for i, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]

                if i == 0:
                    # First row as headers
                    headers = row_data
                else:
                    rows.append(row_data)

            if headers or rows:
                tables.append(
                    TableData(
                        headers=headers if headers else [""] * len(rows[0]) if rows else [],
                        rows=rows,
                        caption=None,
                        page_number=None,
                    )
                )

        return tables

    def _attach_tables_to_sections(
        self, sections: list[DocumentSection], tables: list[TableData]
    ):
        """Attach tables to relevant sections based on content."""
        # Simple heuristic: attach tables to sections that mention them
        # or to BOQ/cost sections
        for table in tables:
            attached = False

            for section in sections:
                if section.section_type in [SectionType.BOQ, SectionType.COST_ESTIMATE]:
                    section.tables.append(table)
                    attached = True
                    break

            # If not attached, add to last section or create new one
            if not attached and sections:
                sections[-1].tables.append(table)

    def _generate_markdown(
        self, sections: list[DocumentSection], tables: list[TableData]
    ) -> str:
        """Generate markdown representation of document."""
        md_parts = []

        for section in sections:
            # Add section header
            md_parts.append(f"## {section.title}")
            md_parts.append("")

            # Add content
            if section.content:
                md_parts.append(section.content)
                md_parts.append("")

            # Add tables
            for table in section.tables:
                md_parts.append(self._table_to_markdown(table))
                md_parts.append("")

        return "\n".join(md_parts)

    def _table_to_markdown(self, table: TableData) -> str:
        """Convert TableData to markdown table."""
        if not table.headers and not table.rows:
            return ""

        lines = []

        # Headers
        if table.headers:
            lines.append("| " + " | ".join(table.headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(table.headers)) + " |")

        # Rows
        for row in table.rows:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    def _detect_section_type(self, title: str) -> SectionType:
        """Detect section type from title."""
        title_lower = title.lower()

        patterns = {
            SectionType.EXECUTIVE_SUMMARY: ["executive summary", "summary", "abstract", "ملخص"],
            SectionType.INTRODUCTION: ["introduction", "background", "مقدمة"],
            SectionType.SCOPE_OF_WORK: ["scope", "scope of work", "نطاق العمل"],
            SectionType.METHODOLOGY: ["methodology", "approach", "method", "المنهجية"],
            SectionType.TECHNICAL_APPROACH: ["technical", "technical approach", "النهج الفني"],
            SectionType.TIMELINE: ["timeline", "schedule", "الجدول الزمني"],
            SectionType.TEAM_COMPOSITION: ["team", "personnel", "staff", "فريق العمل"],
            SectionType.DELIVERABLES: ["deliverables", "outputs", "المخرجات"],
            SectionType.COST_ESTIMATE: ["cost", "budget", "pricing", "التكلفة"],
            SectionType.BOQ: ["boq", "bill of quantities", "جدول الكميات"],
            SectionType.TERMS_AND_CONDITIONS: ["terms", "conditions", "الشروط"],
            SectionType.APPENDIX: ["appendix", "annex", "الملحق"],
            SectionType.REFERENCES: ["references", "bibliography", "المراجع"],
            SectionType.CONCLUSIONS: ["conclusion", "الخلاصة"],
            SectionType.RECOMMENDATIONS: ["recommendation", "التوصيات"],
        }

        for section_type, keywords in patterns.items():
            for keyword in keywords:
                if keyword in title_lower:
                    return section_type

        return SectionType.OTHER
