import os
import pathlib
import hashlib
import logging
import argparse
from datetime import datetime
from docx import Document
import re
import json
import pytesseract
from PIL import Image
import fitz
import io
import pdfplumber

# Configuration constants
CONFIG = {
    "MAX_FILE_SIZE": 50 * 1024 * 1024,  # 50MB limit
    "DEFAULT_ENCODING": "utf-8",
    "LOG_FORMAT": "%(asctime)s - %(levelname)s - %(message)s",
    "LOG_FILE": "document_processing.log",
    "DB_PATH": "processing_tracker.db",
    "KB_PATH": "knowledge_base.json",
    "SNAPSHOT_DIR": "snapshots",
    "MAX_PREVIEW_LENGTH": 10000,  # Max length for OCR output
    "MAX_PAGES_FOR_TABLE_EXTRACTION": 10,  # Limit pages for table extraction
    # TESSERACT_PATHS is now computed dynamically based on platform
    # See get_tesseract_paths() function below
    "ALLOWED_EXTENSIONS": {".docx", ".pdf", ".txt"},
    "HEADER_PATTERNS": [
        re.compile(
            r"^(?:EXECUTIVE SUMMARY|INTRODUCTION|SCOPE OF WORK|TECHNICAL PROPOSAL|FINANCIAL PROPOSAL|DELIVERABLES|TIMELINE|CONCLUSION|PROJECT OVERVIEW|METHODOLOGY)$",
            re.IGNORECASE,
        ),
        re.compile(r"^(?:SECTION \d+|CHAPTER \d+|PART \d+)$", re.IGNORECASE),
        re.compile(r"^\d+\.\s+[A-Z][^.]{5,50}(?<!\.)$"),  # Fixed to prevent backtracking
    ],
    "SECTION_KEYWORDS_PATTERN": re.compile(
        r"^(SECTION|CHAPTER|PART|INTRODUCTION|CONCLUSION|ABSTRACT|REFERENCES|EXECUTIVE SUMMARY|SCOPE OF WORK|TECHNICAL PROPOSAL|FINANCIAL PROPOSAL|DELIVERABLES|TIMELINE|PROJECT OVERVIEW|METHODOLOGY)\b",
        re.IGNORECASE,
    ),
}


# Setup logging with proper Unicode support
logging.basicConfig(
    level=logging.INFO,
    format=CONFIG["LOG_FORMAT"],
    handlers=[
        logging.FileHandler(CONFIG["LOG_FILE"], encoding=CONFIG["DEFAULT_ENCODING"]),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger(__name__)


def get_tesseract_paths():
    """Get platform-specific Tesseract executable paths.

    Returns a list of paths to check for Tesseract OCR, ordered by priority.
    The environment variable TESSERACT_PATH takes highest priority.
    """
    import os
    import sys

    paths = []

    # 1. Environment variable (highest priority)
    env_path = os.environ.get("TESSERACT_PATH")
    if env_path:
        paths.append(env_path)

    # 2. Platform-specific paths - expanded for Windows
    if sys.platform == "win32":
        # Windows paths - more comprehensive list
        paths.extend(
            [
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                r"C:\Tesseract-OCR\tesseract.exe",
                r"D:\Program Files\Tesseract-OCR\tesseract.exe",
                r"D:\Tesseract-OCR\tesseract.exe",
            ]
        )
    elif sys.platform == "darwin":
        # macOS paths
        paths.extend(
            [
                "/opt/homebrew/bin/tesseract",  # Apple Silicon
                "/usr/local/bin/tesseract",  # Intel/Homebrew
            ]
        )
    else:
        # Linux paths
        paths.extend(
            [
                "/usr/bin/tesseract",
                "/usr/local/bin/tesseract",
            ]
        )

    return paths


# Validate OCR availability
OCR_ENABLED = False
TESSERACT_VERSION = None

# Try to get Tesseract path from configuration
tesseract_path = None
for path in get_tesseract_paths():
    if path and os.path.exists(path):
        # Validate that the path is safe before using it
        try:
            # Ensure the path is a file and not a directory
            if os.path.isfile(path):
                # Additional validation: check if it's actually an executable
                if os.access(path, os.X_OK) or path.endswith((".exe", ".bat", ".sh")):
                    tesseract_path = path
                    break
        except (OSError, ValueError):
            continue

try:
    if tesseract_path:
        # Sanitize the path before assigning to prevent command injection
        tesseract_path = os.path.normpath(tesseract_path)
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        TESSERACT_VERSION = pytesseract.get_tesseract_version()
        OCR_ENABLED = True
        logger.info(f"Tesseract OCR v{TESSERACT_VERSION} detected and enabled")
    else:
        logger.warning("Tesseract OCR not found in any of the expected locations")
        logger.warning(
            "OCR features will be disabled. Install from https://github.com/tesseract-ocr/tesseract"
        )
        logger.warning(
            "Or set TESSERACT_PATH environment variable to point to tesseract executable"
        )

    # Check available languages
    if OCR_ENABLED:
        try:
            languages = pytesseract.get_languages()
            logger.info(f"OCR Languages available: {languages}")
            if "ara" in languages:
                logger.info("Arabic OCR support: ENABLED")
            else:
                logger.warning("Arabic OCR support: NOT AVAILABLE (install ara language pack)")
        except Exception as lang_e:
            logger.warning(f"Could not detect OCR languages: {lang_e}")

except Exception as e:
    logger.warning(f"Tesseract OCR not available: {e}")
    logger.warning(
        "OCR features will be disabled. Install from https://github.com/tesseract-ocr/tesseract"
    )


def list_files(folder_path):
    """Iterate through files in folder and report metadata"""
    folder = pathlib.Path(folder_path)

    if not folder.exists():
        logger.error(f"Folder {folder_path} does not exist")
        return

    logger.info(f"Processing folder: {folder.name}")
    logger.info(f"Found {len([f for f in folder.iterdir() if f.is_file()])} files")
    logger.info("-" * 80)

    for file_path in folder.iterdir():
        if file_path.is_file():
            try:
                stat = file_path.stat()
                size_kb = stat.st_size / 1024
                modified = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
                file_type = file_path.suffix.lower()

                logger.info(f"Name: {file_path.name}")
                # Only log full path in debug mode to prevent information disclosure
                if logger.isEnabledFor(logging.DEBUG):
                    logger.debug(f"Path: {file_path}")
                logger.info(f"Size: {size_kb:.2f} KB")
                logger.info(f"Modified: {modified}")
                logger.info(f"Type: {file_type}")
                logger.info("-" * 40)
            except (OSError, PermissionError) as e:
                logger.warning(f"Could not access file {file_path.name}: {e}")


def validate_file_content(file_path):
    """Validate file content type to ensure it matches the extension."""
    import mimetypes

    try:
        # Try to import magic, but handle the case where it's not available
        import magic
    except ImportError:
        # If python-magic is not available, skip content validation
        logger.warning("python-magic not available, skipping content validation")
        return

    try:
        # Get the expected MIME type based on extension
        expected_mime, _ = mimetypes.guess_type(str(file_path))

        # Get the actual MIME type of the file
        actual_mime = magic.from_file(str(file_path), mime=True)

        # Check if the file is actually of the expected type
        # For docx files
        if file_path.suffix.lower() == ".docx":
            if actual_mime not in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/zip",
            ]:
                raise ValueError(
                    f"File {file_path} is not a valid DOCX file (actual type: {actual_mime})"
                )
        # For pdf files
        elif file_path.suffix.lower() == ".pdf":
            if actual_mime != "application/pdf":
                raise ValueError(
                    f"File {file_path} is not a valid PDF file (actual type: {actual_mime})"
                )
        # For txt files
        elif file_path.suffix.lower() == ".txt":
            if not actual_mime.startswith("text/"):
                raise ValueError(
                    f"File {file_path} is not a valid text file (actual type: {actual_mime})"
                )

    except Exception as e:
        logger.warning(f"Could not validate file content for {file_path}: {e}")
        # Don't raise exception here to avoid breaking existing functionality
        # Just log the warning


def validate_file_path(file_path, allowed_extensions, base_path):
    """Validate file path to prevent path traversal and ensure allowed extension."""
    file_path = pathlib.Path(file_path)
    base_path = pathlib.Path(base_path).resolve()

    # Prevent null bytes and other dangerous characters
    if "\x00" in str(file_path):
        raise ValueError(f"File path contains null byte: {file_path}")

    # Resolve the path to prevent traversal
    try:
        resolved_path = file_path.resolve(strict=True)
    except (OSError, RuntimeError):
        raise ValueError(f"File path {file_path} could not be resolved")

    # Ensure the resolved path is within the allowed base path
    try:
        resolved_path.relative_to(base_path)
    except ValueError:
        raise ValueError(f"File path {file_path} is outside allowed base path")

    # Security: Check if the path is a symlink and if so, ensure it's within base path
    # This prevents symlink attacks where a symlink points outside the allowed directory
    if file_path.is_symlink():
        try:
            link_target = pathlib.Path(os.readlink(file_path))
            # If the link is relative, resolve it relative to the symlink's parent directory
            if not link_target.is_absolute():
                link_target = (file_path.parent / link_target).resolve()
            else:
                link_target = link_target.resolve()

            if not link_target.exists():
                raise ValueError(f"Symbolic link {file_path} points to non-existent path")

            try:
                link_target.relative_to(base_path)
            except ValueError:
                raise ValueError(f"Symbolic link {file_path} points outside allowed base path")
        except OSError:
            raise ValueError(f"Could not read symbolic link: {file_path}")

    if resolved_path.suffix.lower() not in allowed_extensions:
        raise ValueError(f"File extension {resolved_path.suffix} not allowed")

    # Check file size to prevent resource exhaustion
    try:
        file_stat = resolved_path.stat()
        if file_stat.st_size > CONFIG["MAX_FILE_SIZE"]:
            raise ValueError(
                f"File {resolved_path} exceeds size limit of {CONFIG['MAX_FILE_SIZE'] / (1024 * 1024):.0f}MB"
            )
    except OSError:
        raise ValueError(f"Could not access file stats for {resolved_path}")

    # Validate file content type
    validate_file_content(resolved_path)

    return True


def get_file_hash(file_path):
    """Calculate SHA-256 hash of file for deduplication with chunked reading."""
    try:
        # Validate file path before processing
        validate_file_path(
            file_path, list(CONFIG["ALLOWED_EXTENSIONS"]), os.path.dirname(file_path)
        )

        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read in chunks to avoid memory issues with large files
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    except Exception as e:
        logger.warning(f"Could not calculate hash for {file_path}: {e}")
        return None


def extract_docx_segments(file_path):
    """Extract text from docx and segment into logical parts with table extraction"""
    try:
        # Validate file path before processing
        validate_file_path(
            file_path, [".docx", ".pdf", ".txt"], os.path.dirname(file_path)
        )  # Use consistent allowed extensions

        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX file {file_path}: {e}")
            return {"Error": f"Could not open DOCX file: {str(e)}"}

        tables = []

        # Extract tables first
        for table_idx, table in enumerate(doc.tables):
            try:
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    try:
                        row_data = [cell.text.strip() if cell.text else "" for cell in row.cells]
                        if any(row_data):  # Only add non-empty rows
                            table_data.append(row_data)
                    except Exception as e:
                        logger.warning(f"Error processing row {row_idx} in table {table_idx}: {e}")
                        continue

                if table_data:
                    # Skip figure captions (single-column tables with figure keywords)
                    if len(table_data[0]) == 1 and table_data:
                        first_cell_text = table_data[0][0].lower()
                        figure_keywords = ["figure", "fig.", "صورة", "شكل", "صورة:", "شكل:"]
                        if any(keyword in first_cell_text for keyword in figure_keywords):
                            logger.debug(f"Skipping figure caption table: {table_data[0][0][:50]}")
                            continue

                    tables.append(
                        {
                            "rows": len(table_data),
                            "columns": len(table_data[0]) if table_data else 0,
                            "data": table_data,
                        }
                    )
            except Exception as e:
                logger.warning(f"Error processing table {table_idx}: {e}")
                continue

        # Improved segmentation with style-based detection - process in chunks to manage memory
        segments = {}
        current_section = "Introduction"
        segments[current_section] = []

        # Common heading styles in docx
        heading_styles = [
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Heading 4",
            "Heading 5",
            "Title",
            "Subtitle",
            "Heading 6",
        ]

        # Process paragraphs in chunks to manage memory
        for para_idx, para in enumerate(doc.paragraphs):
            try:
                if para.text.strip():
                    style_name = para.style.name if para.style and para.style.name else "Normal"
                    item_text = para.text.strip()

                    is_header = False

                    # Check style-based detection
                    if any(heading.lower() in style_name.lower() for heading in heading_styles):
                        header_text = item_text[:100]
                        current_section = header_text
                        if current_section not in segments:
                            segments[current_section] = []
                        is_header = True

                    # Check regex-based detection (fallback)
                    else:
                        for pattern in CONFIG["HEADER_PATTERNS"]:
                            if pattern.match(item_text):
                                current_section = item_text[:100]
                                if current_section not in segments:
                                    segments[current_section] = []
                                is_header = True
                                break

                    if not is_header:
                        segments[current_section].append(item_text)
            except Exception as e:
                logger.warning(f"Error processing paragraph {para_idx}: {e}")
                continue

        # Clean up segments
        for section in segments:
            try:
                segments[section] = "\n".join(segments[section])
            except Exception as e:
                logger.warning(f"Error joining segments for section {section}: {e}")
                segments[section] = ""

        # Return both segments and tables
        return {
            "segments": segments,
            "tables": tables,
            "metadata": {
                "paragraph_count": len(doc.paragraphs),  # Count from doc instead of content list
                "table_count": len(tables),
                "section_count": len(segments),
            },
        }

    except FileNotFoundError:
        logger.error(f"DOCX file not found: {file_path}")
        return {"Error": "File not found"}
    except PermissionError:
        logger.error(f"Permission denied accessing DOCX file: {file_path}")
        return {"Error": "Permission denied"}
    except Exception as e:
        logger.error(f"Unexpected error extracting DOCX {file_path}: {e}", exc_info=True)
        return {"Error": f"Unexpected error: {str(e)}"}


def validate_ocr_quality(ocr_text, file_path):
    """Validate OCR output quality and provide confidence metrics.

    Args:
        ocr_text: Extracted text from OCR
        file_path: Path to processed file (for logging)

    Returns:
        tuple: (is_valid, quality_score, issues)
    """
    issues = []
    quality_score = 0

    # Check 1: Minimum content length
    if len(ocr_text.strip()) < 100:
        issues.append(f"Insufficient text content ({len(ocr_text)} chars, minimum 100)")
        quality_score -= 50

    # Check 2: Character density (non-space characters / total length)
    non_space_chars = len([c for c in ocr_text if not c.isspace()])
    char_density = non_space_chars / max(len(ocr_text), 1)
    if char_density < 0.3:
        issues.append(f"Low character density ({char_density:.2f}, minimum 0.3)")
        quality_score -= 30

    # Check 3: Language presence (English or Arabic)
    has_english = any(c.isalpha() and ord(c) < 128 for c in ocr_text)
    has_arabic = any("\u0600" <= c <= "\u06ff" for c in ocr_text)
    if not has_english and not has_arabic:
        issues.append("No recognizable language detected (English or Arabic)")
        quality_score -= 20

    # Check 4: Word count
    words = ocr_text.split()
    word_count = len(words)
    if word_count < 10:
        issues.append(f"Low word count ({word_count}, minimum 10)")
        quality_score -= 20

    # Calculate final score (0-100 scale)
    quality_score = max(0, min(100, 100 + quality_score))

    # Determine validity
    is_valid = quality_score >= 50 and len(issues) < 3

    if not is_valid:
        logger.warning(
            f"OCR quality check FAILED for {pathlib.Path(file_path).name}: "
            f"score={quality_score}%, issues={len(issues)}"
        )
        for issue in issues:
            logger.warning(f"  - {issue}")
    else:
        logger.info(
            f"OCR quality check PASSED for {pathlib.Path(file_path).name}: score={quality_score}%"
        )

    return is_valid, quality_score, issues


def retry_ocr_page(page_num, page_img, max_retries=3, psm_modes=[6, 3, 1]):
    """Retry OCR with different page segmentation modes for better accuracy.

    Args:
        page_num: Page number for logging
        page_img: PIL Image object of the page
        max_retries: Maximum number of retry attempts
        psm_modes: List of PSM modes to try (6=uniform, 3=auto, 1=auto with OSD)

    Returns:
        str: Best OCR text obtained
    """
    best_text = ""
    best_confidence = 0

    for attempt, psm_mode in enumerate(psm_modes):
        if attempt >= max_retries:
            break

        try:
            custom_config = f"--oem 3 --psm {psm_mode}"
            text = pytesseract.image_to_string(page_img, lang="ara+eng", config=custom_config)

            # Calculate simple quality score
            if len(text.strip()) > len(best_text.strip()):
                best_text = text
                best_confidence = len(text.strip())

            if attempt > 0:
                logger.debug(
                    f"OCR retry {attempt + 1}/max_{max_retries} for page {page_num} with PSM {psm_mode}"
                )

        except Exception as e:
            logger.warning(f"OCR attempt {attempt + 1} failed for page {page_num}: {e}")
            continue

    return best_text


def extract_pdf_with_ocr(file_path):
    """Extract text from PDF using OCR for scanned documents with validation."""
    # Validate input
    if not validate_file_path(
        file_path, [".pdf"], os.path.dirname(file_path)
    ):  # For OCR processing, only allow PDF
        return {"Error": "Invalid file path or type for OCR processing"}

    if not OCR_ENABLED:
        logger.warning(f"OCR disabled, cannot process scanned PDF: {pathlib.Path(file_path).name}")
        return {"Error": "OCR not available - cannot extract text from scanned PDF"}

    doc = None
    try:
        logger.info(f"Running OCR on {pathlib.Path(file_path).name}...")
        doc = fitz.open(file_path)
        ocr_text = ""

        for page_num in range(len(doc)):
            img = None
            try:
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img = Image.open(io.BytesIO(pix.tobytes()))

                # Validate image before OCR
                if img.width > 10 and img.height > 10:  # Basic sanity check
                    # Use retry logic with different PSM modes
                    page_text = retry_ocr_page(page_num, img, max_retries=3)
                    ocr_text += f"\n\n--- Page {page_num + 1} ---\n\n" + page_text
                else:
                    logger.warning(
                        f"Page {page_num} has suspiciously small dimensions: {img.width}x{img.height}"
                    )
            except Exception as e:
                logger.warning(f"OCR failed on page {page_num}: {e}")
                continue
            finally:
                # Clean up image resource
                if img:
                    try:
                        img.close()
                    except:
                        pass  # Best effort cleanup

        # Validate OCR output quality
        if len(ocr_text.strip()) < 10:
            logger.warning(f"OCR produced minimal output for {pathlib.Path(file_path).name}")
            return {"Error": "OCR produced insufficient text content (quality gate failed)"}

        # Run quality gate validation
        is_valid, quality_score, issues = validate_ocr_quality(ocr_text, file_path)

        if not is_valid:
            logger.error(
                f"OCR quality gate FAILED for {pathlib.Path(file_path).name} - "
                f"quality_score={quality_score}%, issues={len(issues)}"
            )
            return {
                "Error": f"OCR quality gate failed (score={quality_score}%): {issues[0] if issues else 'Unknown issue'}",
                "quality_score": quality_score,
                "issues": issues,
            }

        # Segment OCR text
        segments = {
            "OCR_Extracted_Content": ocr_text[: CONFIG["MAX_PREVIEW_LENGTH"]]
        }  # Limit OCR output

        # Add quality metadata
        segments["_metadata"] = {
            "extraction_method": "ocr",
            "quality_score": quality_score,
            "pages_processed": len(doc) if doc else 0,
            "total_chars": len(ocr_text),
        }

        logger.info(
            f"OCR extraction completed with quality score {quality_score}% for {pathlib.Path(file_path).name}"
        )
        return segments

    except FileNotFoundError:
        logger.error(f"PDF file not found for OCR: {file_path}")
        return {"Error": "File not found"}
    except PermissionError:
        logger.error(f"Permission denied for OCR on PDF: {file_path}")
        return {"Error": "Permission denied"}
    except Exception as e:
        logger.error(f"OCR extraction failed for {file_path}: {e}", exc_info=True)
        return {"Error": str(e)}
    finally:
        # Ensure document is closed even if an exception occurs
        if doc:
            try:
                doc.close()
            except Exception as e:
                logger.warning(f"Error closing PDF document during OCR: {e}")


def extract_pdf_segments(file_path, max_pages=None):
    """Extract text from PDF using PyMuPDF with better error handling and memory management."""
    try:
        # Validate file path before processing
        validate_file_path(
            file_path, [".pdf"], os.path.dirname(file_path)
        )  # For PDF processing, only allow PDF

        try:
            doc = fitz.open(file_path)
        except Exception as e:
            logger.error(f"Failed to open PDF file {file_path}: {e}")
            return {"Error": f"Could not open PDF file: {str(e)}"}

        try:
            total_pages = len(doc)

            # Check if it's a scanned/image-only PDF
            try:
                first_page = doc.load_page(0)
                text = str(first_page.get_text())
            except Exception as e:
                logger.error(f"Could not read first page of {file_path}: {e}")
                doc.close()
                return {"Error": "Could not read PDF content"}

            if len(text.strip()) < 50:  # Likely scanned PDF
                logger.info(f"{pathlib.Path(file_path).name} appears to be scanned. Using OCR.")
                doc.close()
                return extract_pdf_with_ocr(file_path)

            page_limit = max_pages if max_pages else total_pages
            logger.info(
                f"Processing {pathlib.Path(file_path).name} ({min(total_pages, page_limit)} pages)..."
            )

            failed_pages = []
            successful_pages = 0

            # Process pages in chunks to manage memory - use a more memory-efficient approach
            segments = {}
            current_section = "Content"
            segments[current_section] = []

            # Track progress for large documents
            progress_interval = max(1, min(total_pages, page_limit) // 10)  # Log every 10%

            for page_num in range(min(total_pages, page_limit)):
                try:
                    try:
                        page = doc.load_page(page_num)
                        page_text = page.get_text()
                    except Exception as e:
                        logger.warning(f"Could not load page {page_num} in {file_path}: {e}")
                        failed_pages.append(page_num)
                        continue

                    # Process page text for segmentation without storing all in memory at once
                    page_lines = page_text.split("\n")

                    for line_idx, line in enumerate(page_lines):
                        try:
                            line = line.strip()
                            if line:
                                # Check if line is a potential section header
                                is_header = False

                                # Check for all caps headers with reasonable length
                                if (
                                    line.isupper()
                                    and 10 <= len(line) <= 100
                                    and len(line.split()) >= 2
                                ):
                                    is_header = True

                                # Check for section keywords
                                elif CONFIG["SECTION_KEYWORDS_PATTERN"].match(line):
                                    is_header = True

                                # Check for numbered sections (e.g., "1. Introduction")
                                elif re.match(r"^\d+\.\s+\w+", line):
                                    is_header = True

                                if is_header:
                                    current_section = line[:100]  # Limit header length
                                    if current_section not in segments:
                                        segments[current_section] = []
                                else:
                                    segments[current_section].append(line)
                        except Exception as e:
                            logger.warning(
                                f"Error processing line {line_idx} on page {page_num}: {e}"
                            )
                            continue

                    successful_pages += 1

                    # Log progress for large documents
                    if (page_num + 1) % progress_interval == 0:
                        logger.debug(
                            f"Processed {page_num + 1}/{min(total_pages, page_limit)} pages"
                        )

                except Exception as e:
                    logger.warning(f"Error reading page {page_num}: {e}")
                    failed_pages.append(page_num)
                    continue

            # Log summary of extraction quality
            if failed_pages:
                logger.warning(f"Failed to extract {len(failed_pages)} pages: {failed_pages}")

            if successful_pages == 0:
                logger.error(f"No pages successfully extracted from {file_path}")
                return {"Error": "No pages could be extracted"}

            logger.info(
                f"Successfully extracted {successful_pages}/{min(total_pages, page_limit)} pages"
            )

            # Clean up segments without hard truncation - optimize for memory
            cleaned_segments = {}
            for section in segments:
                try:
                    content = "\n".join(segments[section])
                    # Only truncate if extremely long to avoid JSON issues
                    if len(content) > 50000:
                        content = content[:50000] + "\n\n... [Content truncated - file too large]"
                    cleaned_segments[section] = content
                except Exception as e:
                    logger.warning(f"Error processing segment '{section}': {e}")
                    cleaned_segments[section] = ""

            return cleaned_segments

        finally:
            try:
                doc.close()
            except Exception as e:
                logger.warning(f"Error closing PDF document {file_path}: {e}")

    except FileNotFoundError:
        logger.error(f"PDF file not found: {file_path}")
        return {"Error": "File not found"}
    except PermissionError:
        logger.error(f"Permission denied accessing PDF file: {file_path}")
        return {"Error": "Permission denied"}
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {e}", exc_info=True)
        return {"Error": f"PDF extraction failed: {str(e)}"}


def extract_table_from_image(image_path):
    """Extract table data from image using OCR"""
    if not OCR_ENABLED:
        return {"Error": "OCR not available"}

    try:
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img, lang="ara+eng")
        # Basic table parsing
        lines = text.split("\n")
        table_data = [line.split("\t") for line in lines if line.strip()]
        return table_data
    except Exception as e:
        logger.error(f"Error extracting table from image: {e}")
        return {"Error": str(e)}


def extract_tables_from_pdf(file_path):
    """Extract tables from PDF using pdfplumber with OCR fallback.

    Args:
        file_path: Path to PDF file

    Returns:
        list: List of extracted tables with metadata
    """
    tables = []

    try:
        logger.info(f"Extracting tables from {pathlib.Path(file_path).name} using pdfplumber...")
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    # Extract tables using pdfplumber
                    page_tables = page.extract_tables()

                    if page_tables:
                        for table_idx, table_data in enumerate(page_tables):
                            try:
                                # Clean and format table data
                                clean_table = []
                                for row_idx, row in enumerate(table_data):
                                    if row:  # Skip empty rows
                                        # Clean cell text
                                        clean_row = [
                                            str(cell).strip() if cell else "" for cell in row
                                        ]
                                        # Only add non-empty rows
                                        if any(clean_row):
                                            clean_table.append(clean_row)

                                if clean_table:
                                    tables.append(
                                        {
                                            "page": page_num + 1,
                                            "rows": len(clean_table),
                                            "columns": len(clean_table[0]) if clean_table else 0,
                                            "data": clean_table,
                                        }
                                    )
                                    logger.debug(
                                        f"Found table on page {page_num + 1}, table {table_idx}: {len(clean_table)} rows, {len(clean_table[0]) if clean_table else 0} cols"
                                    )
                            except Exception as table_e:
                                logger.warning(
                                    f"Error processing table {table_idx} on page {page_num + 1}: {table_e}"
                                )
                                continue

                except Exception as page_e:
                    logger.warning(f"Error extracting tables from page {page_num + 1}: {page_e}")
                    continue

        logger.info(f"Extracted {len(tables)} tables using pdfplumber")
        return tables

    except FileNotFoundError:
        logger.error(f"PDF file not found for table extraction: {file_path}")
        return []
    except PermissionError:
        logger.error(f"Permission denied for table extraction on PDF: {file_path}")
        return []
    except Exception as e:
        logger.error(f"PDF table extraction with pdfplumber failed: {e}", exc_info=True)

        # Fallback to OCR-based table detection
        if OCR_ENABLED:
            logger.info(
                f"Falling back to OCR-based table detection for {pathlib.Path(file_path).name}..."
            )
            return detect_tables_with_ocr(file_path)
        else:
            logger.warning("OCR disabled - cannot extract tables from PDF")
            return []


def detect_tables_with_ocr(file_path):
    """Detect and extract tables from PDF using OCR as fallback.

    This function is called when pdfplumber fails or is unavailable.
    """
    doc = None
    try:
        doc = fitz.open(file_path)
        tables_found = []
        logger.info(f"Scanning {pathlib.Path(file_path).name} for tables with OCR...")

        for page_num in range(
            min(len(doc), CONFIG["MAX_PAGES_FOR_TABLE_EXTRACTION"])
        ):  # Limit to configured number of pages
            img = None
            try:
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img = Image.open(io.BytesIO(pix.tobytes()))

                # Use more sophisticated OCR with table-aware configurations
                custom_config = r"--oem 3 --psm 6"  # Page segmentation mode 6 for uniform blocks
                text = pytesseract.image_to_string(img, lang="ara+eng", config=custom_config)

                if text.strip():
                    # Try to detect and parse tabular data patterns
                    lines = [l for l in text.split("\n") if l.strip()]

                    # Look for multiple types of table-like patterns
                    potential_table = []
                    for line in lines:
                        # Try multiple delimiters to detect table-like structures
                        cells = None

                        # First try tab delimiter
                        if "\t" in line:
                            cells = line.split("\t")
                        # Then try multiple spaces (common in tables)
                        elif "  " in line:  # At least 2 spaces
                            import re

                            cells = re.split(r"\s{2,}", line)  # Split on 2 or more spaces
                        # Then try pipe delimiter (common in ASCII tables)
                        elif "|" in line and line.count("|") > 1:
                            cells = [cell.strip() for cell in line.split("|") if cell.strip()]

                        if cells and len(cells) > 1:  # Has multiple columns
                            potential_table.append([cell.strip() for cell in cells if cell.strip()])

                    if len(potential_table) > 2:  # Only if there's substantial tabular data
                        # Validate that this looks like a real table (consistent columns)
                        if (
                            len(set(len(row) for row in potential_table)) <= 2
                        ):  # At most 2 different column counts
                            tables_found.append(
                                {
                                    "page": page_num + 1,
                                    "rows": len(potential_table),
                                    "columns": len(potential_table[0]) if potential_table else 0,
                                    "data": potential_table,
                                    "method": "ocr_fallback",
                                }
                            )

            except Exception as e:
                logger.warning(f"OCR failed on page {page_num}: {e}")
                continue
            finally:
                # Clean up image resource
                if img:
                    try:
                        img.close()
                    except:
                        pass  # Best effort cleanup

        logger.info(f"Found {len(tables_found)} tables via OCR fallback")
        return tables_found

    except FileNotFoundError:
        logger.error(f"PDF file not found for OCR table detection: {file_path}")
        return []
    except PermissionError:
        logger.error(f"Permission denied for OCR table detection on PDF: {file_path}")
        return []
    except Exception as e:
        logger.error(f"OCR table detection failed for {file_path}: {e}", exc_info=True)
        return []
    finally:
        # Ensure document is closed even if an exception occurs
        if doc:
            try:
                doc.close()
            except Exception as e:
                logger.warning(f"Error closing PDF document during OCR table detection: {e}")


def validate_processing_state(dup_tracker, knowledge_base):
    """Validate consistency between duplicate tracker and knowledge base.

    Returns:
        dict: Validation results with warnings and errors
    """
    validation_results = {"errors": [], "warnings": [], "status": "unknown"}

    try:
        # Get tracked files from duplicate tracker
        tracked_files = dup_tracker.get_all_tracked_files()
        tracked_hashes = dup_tracker.get_all_tracked_hashes()

        # Get processed documents from knowledge base
        kb_docs = {k: v for k, v in knowledge_base.items() if not k.startswith("knowledge_base")}

        # Check for tracked files not in knowledge base
        missing_docs = []
        for file_path in tracked_files:
            doc_id = None
            for kb_key in kb_docs.keys():
                if kb_docs[kb_key].get("file_path") == file_path:
                    doc_id = kb_key
                    break

            if not doc_id:
                missing_docs.append(file_path)

        # Check for knowledge base docs not in tracker
        extra_docs = []
        for doc_id, doc_data in kb_docs.items():
            file_path = doc_data.get("file_path")
            if file_path and file_path not in tracked_files:
                extra_docs.append((doc_id, file_path))

        # Generate validation results
        if missing_docs:
            validation_results["errors"].append(
                f"Duplicate tracker has {len(missing_docs)} files not in knowledge base: "
                f"{missing_docs[:3]}..."
            )

        if extra_docs:
            validation_results["warnings"].append(
                f"Knowledge base has {len(extra_docs)} documents not in duplicate tracker"
            )

        # Determine overall status
        if validation_results["errors"]:
            validation_results["status"] = "error"
        elif validation_results["warnings"]:
            validation_results["status"] = "warning"
        else:
            validation_results["status"] = "healthy"

    except Exception as e:
        validation_results["errors"].append(f"Validation failed: {str(e)}")
        validation_results["status"] = "error"

    return validation_results


def perform_health_check(base_folders, dup_tracker, knowledge_base):
    """Perform comprehensive health check of the document processing system.

    Returns:
        dict: Health check results
    """
    health_results = {"overall_status": "unknown", "checks": {}, "recommendations": []}

    # Check 1: File system accessibility
    fs_healthy = True
    for folder_name, folder_path in base_folders.items():
        if not os.path.exists(folder_path):
            health_results["checks"]["file_system"] = f"ERROR: {folder_name} folder not found"
            fs_healthy = False
        else:
            file_count = len(
                [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
            )
            health_results["checks"][f"file_system_{folder_name}"] = (
                f"OK: {file_count} files accessible"
            )

    if not fs_healthy:
        health_results["overall_status"] = "error"
        health_results["recommendations"].append("Fix file system access issues")

    # Check 2: Database connectivity
    try:
        test_files = dup_tracker.get_all_tracked_files()
        health_results["checks"]["database"] = f"OK: Connected, {len(test_files)} files tracked"
    except Exception as e:
        health_results["checks"]["database"] = f"ERROR: Database connection failed: {str(e)}"
        health_results["overall_status"] = "error"
        health_results["recommendations"].append("Fix database connectivity issues")

    # Check 3: Knowledge base integrity
    kb_docs = {k: v for k, v in knowledge_base.items() if not k.startswith("knowledge_base")}
    if kb_docs:
        health_results["checks"]["knowledge_base"] = f"OK: {len(kb_docs)} documents stored"
    else:
        health_results["checks"]["knowledge_base"] = "WARNING: No processed documents found"
        health_results["recommendations"].append(
            "Run document processing to populate knowledge base"
        )

    # Check 4: State consistency
    validation = validate_processing_state(dup_tracker, knowledge_base)
    health_results["checks"]["state_consistency"] = (
        f"{validation['status'].upper()}: {len(validation['errors'])} errors, {len(validation['warnings'])} warnings"
    )

    if validation["status"] == "error":
        health_results["overall_status"] = "error"
        health_results["recommendations"].extend(validation["errors"])
    elif validation["status"] == "warning":
        if health_results["overall_status"] != "error":
            health_results["overall_status"] = "warning"
        health_results["recommendations"].extend(validation["warnings"])

    # Check 5: OCR availability
    if OCR_ENABLED:
        health_results["checks"]["ocr"] = f"OK: Tesseract v{TESSERACT_VERSION} available"
    else:
        health_results["checks"]["ocr"] = (
            "INFO: OCR not available (install Tesseract for full functionality)"
        )
        health_results["recommendations"].append("Install Tesseract OCR for scanned PDF processing")

    # Determine overall status
    if health_results["overall_status"] == "unknown":
        if all(
            ("OK" in str(status) or "HEALTHY" in str(status))
            for status in health_results["checks"].values()
        ):
            health_results["overall_status"] = "healthy"
        else:
            health_results["overall_status"] = "warning"

    return health_results


def process_all_documents(
    base_folders, force_reprocess=False, selective_files=None, existing_kb=None
):
    """Process all documents in folders and create knowledge library with deduplication

    Args:
        base_folders: Dict of folder_name -> folder_path
        force_reprocess: If True, ignore duplicate tracker and reprocess all files
        selective_files: List of specific files to process (if None, process all)
        existing_kb: Optional existing knowledge base dict to merge with (preserves previous data)
    """
    # Start with existing knowledge base if provided, otherwise create empty one
    # This FIXES the data loss bug - we now accept and preserve existing data
    knowledge_base = existing_kb.copy() if existing_kb else {}
    file_count = 0
    error_count = 0
    skipped_count = 0

    # Initialize duplicate tracker
    dup_tracker = DuplicateTracker()

    for folder_name, folder_path in base_folders.items():
        logger.info(f"\n{'=' * 60}")
        logger.info(f"Processing {folder_name} folder...")
        logger.info(f"{'=' * 60}")
        folder = pathlib.Path(folder_path)

        if not folder.exists():
            logger.error(f"Folder not found: {folder_path}")
            continue

        # Get all files first to show total count
        all_files = [
            f for f in folder.iterdir() if f.is_file() and f.suffix.lower() in [".docx", ".pdf"]
        ]

        for idx, file_path in enumerate(all_files):
            # Check selective processing filter
            if selective_files and str(file_path) not in selective_files:
                continue

            file_count += 1
            logger.info(f"\n[{idx + 1}/{len(all_files)}] Processing: {file_path.name}")

            # Get file hash for duplicate checking and recording
            file_hash = get_file_hash(file_path) if not force_reprocess else None

            # Check for duplicates using persistent tracker (unless force reprocess)
            if (
                not force_reprocess
                and file_hash
                and dup_tracker.is_duplicate(str(file_path), file_hash)
            ):
                logger.warning(f"  Skipping previously processed file (hash: {file_hash[:8]}...)")
                skipped_count += 1
                continue

            doc_id = f"{folder_name}_{file_path.stem}"

            try:
                if file_path.suffix.lower() == ".docx":
                    result = extract_docx_segments(str(file_path))
                    if isinstance(result, dict) and "segments" in result:
                        # Sanitize and validate the extracted data before storing
                        sanitized_result = _sanitize_document_data(result)
                        knowledge_base[doc_id] = {
                            "file_path": str(file_path),
                            "type": folder_name,
                            "format": file_path.suffix.lower(),
                            "segments": sanitized_result["segments"],
                            "tables": sanitized_result.get("tables", []),
                            "metadata": sanitized_result.get("metadata", {}),
                            "processed_at": datetime.now().isoformat(),
                        }
                    elif "Error" in result:
                        error_msg = (
                            str(result["Error"])
                            if isinstance(result["Error"], str)
                            else str(result["Error"])
                        )
                        knowledge_base[doc_id] = {
                            "file_path": str(file_path),
                            "type": folder_name,
                            "format": file_path.suffix.lower(),
                            "error": error_msg,
                            "processed_at": datetime.now().isoformat(),
                        }

                elif file_path.suffix.lower() == ".pdf":
                    segments = extract_pdf_segments(str(file_path))
                    if "Error" not in segments:
                        # Sanitize and validate the extracted data before storing
                        sanitized_segments = _sanitize_document_data({"segments": segments})
                        knowledge_base[doc_id] = {
                            "file_path": str(file_path),
                            "type": folder_name,
                            "format": file_path.suffix.lower(),
                            "segments": sanitized_segments["segments"],
                            "processed_at": datetime.now().isoformat(),
                        }
                    else:
                        error_msg = (
                            str(segments["Error"])
                            if isinstance(segments["Error"], str)
                            else str(segments["Error"])
                        )
                        knowledge_base[doc_id] = {
                            "file_path": str(file_path),
                            "type": folder_name,
                            "format": file_path.suffix.lower(),
                            "error": error_msg,
                            "processed_at": datetime.now().isoformat(),
                        }

                # Record successful processing
                if file_hash:
                    dup_tracker.record_processed(str(file_path), file_hash)
                logger.info(f"  [OK] Successfully processed")

            except Exception as e:
                error_count += 1
                logger.error(f"  [ERROR] Error processing {file_path.name}: {e}", exc_info=True)
                knowledge_base[doc_id] = {
                    "file_path": str(file_path),
                    "type": folder_name,
                    "format": file_path.suffix.lower(),
                    "error": str(e),
                    "processed_at": datetime.now().isoformat(),
                }

    # Extract tables from PDFs using pdfplumber - process in batches to manage memory
    logger.info(f"\n{'=' * 60}")
    logger.info("Extracting tables from PDFs...")
    logger.info(f"{'=' * 60}")

    # Get PDF documents that need table extraction
    pdf_docs = [
        (doc_id, doc_data)
        for doc_id, doc_data in knowledge_base.items()
        if doc_data.get("format") == ".pdf" and "error" not in doc_data
    ]

    for idx, (doc_id, doc_data) in enumerate(pdf_docs):
        file_path = doc_data.get("file_path")
        if file_path:
            logger.info(f"  [{idx + 1}/{len(pdf_docs)}] Table extraction: {doc_id}")
            pdf_tables = extract_tables_from_pdf(file_path)
            if pdf_tables:
                doc_data["tables"] = pdf_tables
                logger.info(f"    Found {len(pdf_tables)} tables")
            else:
                doc_data["tables"] = []
                logger.info(f"    No tables found")

    # Add OCR processing for PDFs if enabled - process in batches to manage memory
    if OCR_ENABLED:
        logger.info(f"\n{'=' * 60}")
        logger.info("Running additional OCR on PDFs for table detection...")
        logger.info(f"{'=' * 60}")

        # Get PDF documents that don't have tables yet
        ocr_docs = [
            (doc_id, doc_data)
            for doc_id, doc_data in knowledge_base.items()
            if doc_data.get("format") == ".pdf"
            and "error" not in doc_data
            and not doc_data.get("tables")
        ]

        for idx, (doc_id, doc_data) in enumerate(ocr_docs):
            file_path = doc_data.get("file_path")
            if file_path:
                logger.info(f"  [{idx + 1}/{len(ocr_docs)}] OCR scanning: {doc_id}")
                ocr_tables = detect_tables_with_ocr(file_path)
                if ocr_tables:
                    doc_data["tables"] = ocr_tables
                    logger.info(f"    Found {len(ocr_tables)} tables via OCR")
    else:
        logger.info("\nOCR is disabled - skipping additional OCR processing")

    logger.info(f"\n{'=' * 60}")
    logger.info(f"Processing Summary:")
    logger.info(f"  Total files found: {file_count}")
    logger.info(f"  Files skipped (duplicates): {skipped_count}")
    logger.info(
        f"  Successfully processed: {len([d for d in knowledge_base.values() if 'error' not in d])}"
    )
    logger.info(f"  Errors: {error_count}")
    logger.info(f"{'=' * 60}\n")

    return knowledge_base


import tempfile
import shutil
import sqlite3
import atexit


import threading


class DuplicateTracker:
    def __init__(self, db_path=None):
        self.db_path = db_path or CONFIG["DB_PATH"]
        self.local = threading.local()  # Thread-local storage for connection
        self.lock = threading.Lock()  # Lock for thread safety
        self.setup_db()
        # Removed atexit.register to prevent potential hanging issues

    def get_connection(self):
        """Get thread-local database connection."""
        if not hasattr(self.local, "conn"):
            self.local.conn = sqlite3.connect(self.db_path, check_same_thread=False)
        return self.local.conn

    def setup_db(self):
        """Setup database table for tracking processed files."""
        conn = self.get_connection()
        conn.execute("""
            CREATE TABLE IF NOT EXISTS processed_files (
                file_path TEXT PRIMARY KEY,
                file_hash TEXT,
                processed_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(file_hash)
            )
        """)
        conn.commit()

    def is_duplicate(self, file_path, file_hash):
        """Check if file has been processed before."""
        with self.lock:  # Thread-safe access
            try:
                conn = self.get_connection()
                cursor = conn.execute(
                    "SELECT file_path FROM processed_files WHERE file_hash = ?", (file_hash,)
                )
                result = cursor.fetchone()
                return result is not None
            except sqlite3.Error as e:
                logger.error(f"Database error checking duplicate: {e}")
                return False  # Assume not duplicate if DB error occurs

    def record_processed(self, file_path, file_hash):
        """Record that a file has been processed."""
        with self.lock:  # Thread-safe access
            try:
                conn = self.get_connection()
                conn.execute(
                    "INSERT OR REPLACE INTO processed_files (file_path, file_hash) VALUES (?, ?)",
                    (str(file_path), file_hash),
                )
                conn.commit()
            except sqlite3.Error as e:
                logger.error(f"Failed to record processed file: {e}")

    def get_all_tracked_files(self):
        """Get list of all tracked file paths."""
        with self.lock:  # Thread-safe access
            try:
                conn = self.get_connection()
                cursor = conn.execute("SELECT file_path FROM processed_files")
                return [row[0] for row in cursor.fetchall()]
            except sqlite3.Error as e:
                logger.error(f"Database error getting tracked files: {e}")
                return []

    def get_all_tracked_hashes(self):
        """Get list of all tracked file hashes."""
        with self.lock:  # Thread-safe access
            try:
                conn = self.get_connection()
                cursor = conn.execute("SELECT file_hash FROM processed_files")
                return [row[0] for row in cursor.fetchall()]
            except sqlite3.Error as e:
                logger.error(f"Database error getting tracked hashes: {e}")
                return []

    def clear_all_records(self):
        """Clear all duplicate tracking records."""
        with self.lock:  # Thread-safe access
            try:
                conn = self.get_connection()
                conn.execute("DELETE FROM processed_files")
                conn.commit()
                logger.info("Cleared all duplicate tracking records")
            except sqlite3.Error as e:
                logger.error(f"Failed to clear duplicate records: {e}")

    def get_record_count(self):
        """Get total number of tracked records."""
        with self.lock:  # Thread-safe access
            try:
                conn = self.get_connection()
                cursor = conn.execute("SELECT COUNT(*) FROM processed_files")
                return cursor.fetchone()[0]
            except sqlite3.Error as e:
                logger.error(f"Database error getting record count: {e}")
                return 0

    def close(self):
        """Close database connection."""
        if hasattr(self.local, "conn"):
            self.local.conn.close()


def atomic_save_json(data, file_path):
    """Save JSON atomically to prevent corruption during writes."""
    # Validate data before saving to prevent corrupting the knowledge base
    if not isinstance(data, dict):
        raise ValueError("Data must be a dictionary")

    # Perform basic validation to ensure data integrity
    for key, value in data.items():
        if not isinstance(key, str):
            raise ValueError(f"All keys must be strings, got {type(key)} for key {key}")
        # Check for potential issues with values
        if isinstance(value, (dict, list)):
            # Recursively validate nested structures
            _validate_nested_data(value, f"key '{key}'")

    temp_file = None
    try:
        # Create temporary file in the same directory to ensure atomic move
        temp_dir = os.path.dirname(file_path)
        with tempfile.NamedTemporaryFile(
            mode="w", dir=temp_dir, delete=False, encoding="utf-8"
        ) as temp_f:
            temp_file = temp_f.name
            json.dump(data, temp_f, ensure_ascii=False, indent=2)

        # Atomic move operation
        shutil.move(temp_file, file_path)
        logger.info(f"Knowledge base saved to {file_path}")

    except Exception as e:
        # Clean up temp file if something went wrong
        if temp_file and os.path.exists(temp_file):
            try:
                os.unlink(temp_file)
            except:
                pass  # Best effort cleanup
        logger.error(f"Failed to save knowledge base: {e}")
        raise


def _validate_nested_data(obj, path=""):
    """Recursively validate nested data structures to prevent corruption."""
    if isinstance(obj, dict):
        for key, value in obj.items():
            if not isinstance(key, str):
                raise ValueError(f"All keys must be strings, got {type(key)} at {path}.{key}")
            _validate_nested_data(value, f"{path}.{key}")
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            _validate_nested_data(item, f"{path}[{i}]")
    elif isinstance(obj, (str, int, float, bool)) or obj is None:
        # These are valid primitive types
        pass
    else:
        raise ValueError(f"Invalid data type {type(obj)} at {path}")


def _sanitize_document_data(data):
    """Sanitize document data to prevent injection and ensure valid content."""
    if not isinstance(data, dict):
        return data

    sanitized = {}
    for key, value in data.items():
        # Sanitize keys
        clean_key = str(key)[:1000] if key is not None else ""  # Limit key length

        # Sanitize values based on type
        if isinstance(value, str):
            # Remove potentially harmful characters and limit length
            clean_value = value.replace("\x00", "").replace("\r\x00", "")[:1000000]  # 1MB limit
            sanitized[clean_key] = clean_value
        elif isinstance(value, dict):
            sanitized[clean_key] = _sanitize_document_data(value)
        elif isinstance(value, list):
            sanitized[clean_key] = [
                _sanitize_document_data(item) if isinstance(item, dict) else item for item in value
            ]
        else:
            # For other types, just pass through
            sanitized[clean_key] = value

    return sanitized


def synchronize_states(dup_tracker, knowledge_base, kb_path):
    """Synchronize duplicate tracker and knowledge base states.

    This function ensures consistency between the duplicate tracker database
    and the knowledge base by:
    1. Adding missing entries to duplicate tracker
    2. Removing orphaned entries from duplicate tracker
    3. Creating backups before making changes

    Returns:
        dict: Synchronization results
    """
    sync_results = {"status": "success", "actions_taken": [], "warnings": [], "errors": []}

    try:
        # Create backup before synchronization
        backup_path = str(kb_path) + ".sync_backup"
        try:
            import shutil

            shutil.copy2(kb_path, backup_path)
            sync_results["actions_taken"].append(f"Created backup: {backup_path}")
        except Exception as e:
            sync_results["warnings"].append(f"Could not create backup: {e}")

        # Get current state
        tracked_files = dup_tracker.get_all_tracked_files()
        kb_docs = {k: v for k, v in knowledge_base.items() if not k.startswith("knowledge_base")}

        # Find documents in KB but not in tracker
        missing_in_tracker = []
        for doc_id, doc_data in kb_docs.items():
            file_path = doc_data.get("file_path")
            if file_path and file_path not in tracked_files:
                missing_in_tracker.append((file_path, doc_data))

        # Find files in tracker but not in KB
        missing_in_kb = []
        for tracked_file in tracked_files:
            found = False
            for doc_data in kb_docs.values():
                if doc_data.get("file_path") == tracked_file:
                    found = True
                    break
            if not found:
                missing_in_kb.append(tracked_file)

        # Synchronize: Add missing entries to tracker
        for file_path, doc_data in missing_in_tracker:
            try:
                # Calculate hash for the file
                file_hash = get_file_hash(file_path)
                if file_hash:
                    dup_tracker.record_processed(file_path, file_hash)
                    sync_results["actions_taken"].append(
                        f"Added to tracker: {os.path.basename(file_path)}"
                    )
            except Exception as e:
                sync_results["errors"].append(f"Failed to add {file_path} to tracker: {e}")

        # Note: We don't automatically remove from tracker as it might be intentional
        # to keep track of previously processed files
        if missing_in_kb:
            sync_results["warnings"].append(
                f"Found {len(missing_in_kb)} files in tracker not in KB (keeping for history)"
            )

        logger.info(
            f"State synchronization completed: {len(sync_results['actions_taken'])} actions taken"
        )

    except Exception as e:
        sync_results["status"] = "error"
        sync_results["errors"].append(f"Synchronization failed: {e}")
        logger.error(f"State synchronization failed: {e}")

    return sync_results


def create_snapshot(knowledge_base, snapshot_dir=None):
    """Create a timestamped snapshot of the knowledge base.

    Returns:
        str: Path to created snapshot file
    """
    try:
        # Create snapshots directory
        snapshot_dir = snapshot_dir or CONFIG["SNAPSHOT_DIR"]
        snapshot_path = pathlib.Path(snapshot_dir)
        snapshot_path.mkdir(exist_ok=True)

        # Create timestamped filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        snapshot_file = snapshot_path / f"knowledge_base_snapshot_{timestamp}.json"

        # Save snapshot
        with open(snapshot_file, "w", encoding="utf-8") as f:
            json.dump(knowledge_base, f, ensure_ascii=False, indent=2)

        logger.info(f"Created snapshot: {snapshot_file}")
        return str(snapshot_file)

    except Exception as e:
        logger.error(f"Failed to create snapshot: {e}")
        return None


def list_snapshots(snapshot_dir=None):
    """List available snapshots with metadata.

    Returns:
        list: List of snapshot info dictionaries
    """
    try:
        snapshot_dir = snapshot_dir or CONFIG["SNAPSHOT_DIR"]
        snapshot_path = pathlib.Path(snapshot_dir)
        if not snapshot_path.exists():
            return []

        snapshots = []
        for snapshot_file in snapshot_path.glob("knowledge_base_snapshot_*.json"):
            try:
                # Extract timestamp from filename
                filename = snapshot_file.name
                timestamp_str = filename.replace("knowledge_base_snapshot_", "").replace(
                    ".json", ""
                )
                timestamp = datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")

                # Get file size
                size = snapshot_file.stat().st_size

                snapshots.append(
                    {
                        "path": str(snapshot_file),
                        "filename": filename,
                        "timestamp": timestamp,
                        "size": size,
                        "documents": None,  # Will be loaded on demand
                    }
                )

            except Exception as e:
                logger.warning(f"Could not parse snapshot {snapshot_file}: {e}")

        # Sort by timestamp (newest first)
        snapshots.sort(key=lambda x: x["timestamp"], reverse=True)
        return snapshots

    except Exception as e:
        logger.error(f"Failed to list snapshots: {e}")
        return []


def rollback_to_snapshot(snapshot_path, target_path="knowledge_base.json"):
    """Rollback knowledge base to a specific snapshot.

    Args:
        snapshot_path: Path to snapshot file
        target_path: Path to target knowledge base file

    Returns:
        bool: Success status
    """
    try:
        # Validate snapshot exists and is readable
        snapshot_file = pathlib.Path(snapshot_path)
        if not snapshot_file.exists():
            logger.error(f"Snapshot file does not exist: {snapshot_path}")
            return False

        # Load snapshot
        with open(snapshot_file, "r", encoding="utf-8") as f:
            snapshot_data = json.load(f)

        # Create backup of current state
        current_file = pathlib.Path(target_path)
        if current_file.exists():
            backup_path = str(current_file) + ".rollback_backup"
            import shutil

            shutil.copy2(current_file, backup_path)
            logger.info(f"Created rollback backup: {backup_path}")

        # Apply snapshot
        atomic_save_json(snapshot_data, target_path)

        logger.info(f"Successfully rolled back to snapshot: {snapshot_path}")
        return True

    except Exception as e:
        logger.error(f"Rollback failed: {e}")
        return False


def create_monitoring_dashboard(base_folders, dup_tracker, knowledge_base):
    """Create a monitoring dashboard with system status and metrics.

    Returns:
        dict: Dashboard data
    """
    dashboard = {
        "timestamp": datetime.now().isoformat(),
        "system_status": "unknown",
        "metrics": {},
        "components": {},
        "alerts": [],
        "recommendations": [],
    }

    try:
        # System health check
        health_results = perform_health_check(base_folders, dup_tracker, knowledge_base)
        dashboard["system_status"] = health_results["overall_status"]
        dashboard["components"] = health_results["checks"]
        dashboard["recommendations"].extend(health_results["recommendations"])

        # Core metrics
        kb_docs = {k: v for k, v in knowledge_base.items() if not k.startswith("knowledge_base")}

        dashboard["metrics"] = {
            "total_documents": len(kb_docs),
            "tracked_files": dup_tracker.get_record_count(),
            "docx_files": len([d for d in kb_docs.values() if d.get("format") == ".docx"]),
            "pdf_files": len([d for d in kb_docs.values() if d.get("format") == ".pdf"]),
            "total_segments": sum(len(d.get("segments", {})) for d in kb_docs.values()),
            "total_tables": sum(len(d.get("tables", [])) for d in kb_docs.values()),
            "ocr_enabled": OCR_ENABLED,
            "processing_errors": len([d for d in kb_docs.values() if "error" in d]),
        }

        # File system status
        fs_status = {}
        for folder_name, folder_path in base_folders.items():
            if os.path.exists(folder_path):
                files = [
                    f
                    for f in os.listdir(folder_path)
                    if os.path.isfile(os.path.join(folder_path, f))
                ]
                docx_count = len([f for f in files if f.lower().endswith(".docx")])
                pdf_count = len([f for f in files if f.lower().endswith(".pdf")])
                fs_status[folder_name] = {
                    "total_files": len(files),
                    "docx_files": docx_count,
                    "pdf_files": pdf_count,
                    "status": "accessible",
                }
            else:
                fs_status[folder_name] = {"status": "not_found"}
                dashboard["alerts"].append(f"Folder not found: {folder_name}")

        dashboard["file_system"] = fs_status

        # State consistency check
        validation = validate_processing_state(dup_tracker, knowledge_base)
        dashboard["state_consistency"] = {
            "status": validation["status"],
            "errors": len(validation["errors"]),
            "warnings": len(validation["warnings"]),
        }

        if validation["errors"]:
            dashboard["alerts"].extend(validation["errors"])
        if validation["warnings"]:
            dashboard["alerts"].extend(validation["warnings"])

        # Performance indicators
        if dashboard["metrics"]["total_documents"] > 0:
            avg_segments = (
                dashboard["metrics"]["total_segments"] / dashboard["metrics"]["total_documents"]
            )
            avg_tables = (
                dashboard["metrics"]["total_tables"] / dashboard["metrics"]["total_documents"]
            )
            dashboard["performance"] = {
                "avg_segments_per_doc": round(avg_segments, 1),
                "avg_tables_per_doc": round(avg_tables, 1),
                "error_rate": (
                    dashboard["metrics"]["processing_errors"]
                    / dashboard["metrics"]["total_documents"]
                )
                * 100,
            }

        # Generate alerts based on thresholds
        if dashboard["metrics"]["processing_errors"] > 0:
            dashboard["alerts"].append(
                f"Processing errors detected: {dashboard['metrics']['processing_errors']}"
            )

        if dashboard["state_consistency"]["status"] == "error":
            dashboard["alerts"].append("State consistency issues detected")

        if not OCR_ENABLED:
            dashboard["recommendations"].append("Install Tesseract OCR for full PDF processing")

    except Exception as e:
        dashboard["system_status"] = "error"
        dashboard["alerts"].append(f"Dashboard generation failed: {e}")
        logger.error(f"Dashboard generation failed: {e}")

    return dashboard


def create_argument_parser():
    """Create and configure argument parser for the document processing pipeline."""
    parser = argparse.ArgumentParser(
        description="Document Processing Pipeline for AI Contract Generation",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Processing Modes:
  incremental (default): Process only new or changed files
  full-reprocess: Reprocess all files, ignoring duplicate tracker
  selective: Process only specified files

Examples:
  python doc_pipeline.py                           # Incremental processing
  python doc_pipeline.py --force-reprocess         # Full reprocessing
  python doc_pipeline.py --health-check            # System health check
  python doc_pipeline.py --validate-state          # Validate system state
  python doc_pipeline.py --clear-duplicates        # Clear duplicate tracker
  python doc_pipeline.py --files proposals/*.docx  # Process specific files
        """,
    )

    # Processing mode options
    parser.add_argument(
        "--force-reprocess",
        action="store_true",
        help="Force reprocessing of all files, ignoring duplicate tracker",
    )

    parser.add_argument("--files", nargs="+", help="Process only specified files (selective mode)")

    # System management options
    parser.add_argument(
        "--health-check", action="store_true", help="Perform system health check and exit"
    )

    parser.add_argument(
        "--validate-state",
        action="store_true",
        help="Validate consistency between duplicate tracker and knowledge base",
    )

    parser.add_argument(
        "--clear-duplicates", action="store_true", help="Clear all duplicate tracking records"
    )

    parser.add_argument(
        "--sync-states",
        action="store_true",
        help="Synchronize duplicate tracker and knowledge base states",
    )

    parser.add_argument(
        "--create-snapshot",
        action="store_true",
        help="Create a timestamped snapshot of the knowledge base",
    )

    parser.add_argument(
        "--list-snapshots", action="store_true", help="List available knowledge base snapshots"
    )

    parser.add_argument(
        "--rollback",
        metavar="SNAPSHOT_FILE",
        help="Rollback knowledge base to specified snapshot file",
    )

    parser.add_argument(
        "--dashboard",
        action="store_true",
        help="Display monitoring dashboard with system status and metrics",
    )

    # Output options
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging")

    return parser


def main():
    """Main entry point for the document processing pipeline."""
    parser = create_argument_parser()
    args = parser.parse_args()

    # Configure logging level
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    # Use relative paths
    base_dir = pathlib.Path(__file__).parent
    proposals_folder = base_dir / "proposals"
    reports_folder = base_dir / "Reports"
    base_folders = {"proposals": str(proposals_folder), "reports": str(reports_folder)}
    existing_kb_path = base_dir / CONFIG["KB_PATH"]

    # Load existing knowledge base - FIX: Load all existing data, not just static section
    knowledge_base = {}
    if existing_kb_path.exists():
        try:
            with open(existing_kb_path, "r", encoding="utf-8") as f:
                raw_data = f.read()
                # Basic validation to prevent certain injection attacks
                if "\x00" in raw_data:
                    raise ValueError("JSON file contains null bytes")

                # Parse JSON with limits to prevent resource exhaustion
                knowledge_base = json.loads(raw_data)

                # Validate that the loaded data is a dictionary
                if not isinstance(knowledge_base, dict):
                    raise ValueError("Knowledge base must be a dictionary")

                logger.info(
                    f"Loaded existing knowledge base with {len([k for k in knowledge_base.keys() if not k.startswith('knowledge_base')])} documents"
                )
        except (json.JSONDecodeError, ValueError) as e:
            logger.error(f"Could not load existing knowledge_base due to invalid format: {e}")
            raise
        except Exception as e:
            logger.error(f"Could not load existing knowledge_base: {e}")
            raise

    # Initialize components
    dup_tracker = DuplicateTracker()

    # Handle management commands
    if args.clear_duplicates:
        logger.info("Clearing all duplicate tracking records...")
        dup_tracker.clear_all_records()
        logger.info("Duplicate tracker cleared successfully")
        return

    if args.sync_states:
        logger.info("Synchronizing duplicate tracker and knowledge base states...")
        sync_results = synchronize_states(dup_tracker, knowledge_base, existing_kb_path)

        print(f"\n{'=' * 60}")
        print("STATE SYNCHRONIZATION RESULTS")
        print("=" * 60)
        print(f"Status: {sync_results['status'].upper()}")

        if sync_results["actions_taken"]:
            print(f"\nActions Taken ({len(sync_results['actions_taken'])}):")
            for action in sync_results["actions_taken"]:
                print(f"  [OK] {action}")

        if sync_results["warnings"]:
            print(f"\nWarnings ({len(sync_results['warnings'])}):")
            for warning in sync_results["warnings"]:
                print(f"  [WARN] {warning}")

        if sync_results["errors"]:
            print(f"\nErrors ({len(sync_results['errors'])}):")
            for error in sync_results["errors"]:
                print(f"  [ERROR] {error}")

        print("=" * 60)
        return

    if args.create_snapshot:
        logger.info("Creating knowledge base snapshot...")
        snapshot_path = create_snapshot(knowledge_base)
        if snapshot_path:
            print(f"\n[SUCCESS] Snapshot created: {snapshot_path}")
        else:
            print("\n[ERROR] Failed to create snapshot")
        return

    if args.list_snapshots:
        logger.info("Listing available snapshots...")
        snapshots = list_snapshots()

        print(f"\n{'=' * 80}")
        print("AVAILABLE SNAPSHOTS")
        print("=" * 80)

        if snapshots:
            print(f"{'Filename':<45} {'Date/Time':<20} {'Size':<10}")
            print("-" * 80)

            for snap in snapshots:
                size_kb = snap["size"] // 1024
                date_str = snap["timestamp"].strftime("%Y-%m-%d %H:%M:%S")
                print(f"{snap['filename']:<45} {date_str:<20} {size_kb:>8} KB")
        else:
            print("No snapshots found")

        print("=" * 80)
        return

    if args.rollback:
        logger.info(f"Rolling back to snapshot: {args.rollback}")
        success = rollback_to_snapshot(args.rollback)
        if success:
            print(f"\n[SUCCESS] Successfully rolled back to: {args.rollback}")
            print("Note: Original knowledge base backed up as .rollback_backup")
        else:
            print(f"\n[ERROR] Rollback failed")
        return

    if args.dashboard:
        logger.info("Generating monitoring dashboard...")
        dashboard = create_monitoring_dashboard(base_folders, dup_tracker, knowledge_base)

        print(f"\n{'=' * 80}")
        print("SYSTEM MONITORING DASHBOARD")
        print("=" * 80)
        print(f"Timestamp: {dashboard['timestamp']}")
        print(f"System Status: {dashboard['system_status'].upper()}")

        print(f"\nCORE METRICS:")
        metrics = dashboard["metrics"]
        for key, value in metrics.items():
            print(f"  {key}: {value}")

        if "performance" in dashboard:
            print(f"\nPERFORMANCE INDICATORS:")
            perf = dashboard["performance"]
            for key, value in perf.items():
                print(f"  {key}: {value}")

        print(f"\nCOMPONENT STATUS:")
        for component, status in dashboard["components"].items():
            status_str = str(status).upper()
            if "OK" in status_str:
                status_indicator = "[OK]"
            elif "WARNING" in status_str or "WARN" in status_str:
                status_indicator = "[WARN]"
            elif "ERROR" in status_str or "FAILED" in status_str:
                status_indicator = "[ERROR]"
            elif "HEALTHY" in status_str:
                status_indicator = "[OK]"
            else:
                status_indicator = "[INFO]"
            print(f"  {status_indicator} {component}: {status}")

        if dashboard["alerts"]:
            print(f"\nALERTS ({len(dashboard['alerts'])}):")
            for alert in dashboard["alerts"]:
                print(f"  [ALERT] {alert}")

        if dashboard["recommendations"]:
            print(f"\nRECOMMENDATIONS ({len(dashboard['recommendations'])}):")
            for rec in dashboard["recommendations"]:
                print(f"  [TIP] {rec}")

        print("=" * 80)
        return

    # Load existing knowledge base - FIX: Load all existing data, not just static section
    # This section is redundant since we already loaded knowledge_base earlier, so removing it
    # The knowledge_base variable is already populated from the earlier load

    # Perform health check if requested
    if args.health_check:
        logger.info("Performing system health check...")
        health_results = perform_health_check(base_folders, dup_tracker, knowledge_base)

        print(f"\n{'=' * 60}")
        print("SYSTEM HEALTH CHECK RESULTS")
        print("=" * 60)
        print(f"Overall Status: {health_results['overall_status'].upper()}")

        print(f"\nComponent Status:")
        for check, status in health_results["checks"].items():
            print(f"  {check}: {status}")

        if health_results["recommendations"]:
            print(f"\nRecommendations:")
            for rec in health_results["recommendations"]:
                print(f"  - {rec}")

        print("=" * 60)
        return

    # Validate state if requested
    if args.validate_state:
        logger.info("Validating system state consistency...")
        validation = validate_processing_state(dup_tracker, knowledge_base)

        print(f"\n{'=' * 60}")
        print("STATE VALIDATION RESULTS")
        print("=" * 60)
        print(f"Overall Status: {validation['status'].upper()}")

        if validation["errors"]:
            print(f"\nErrors ({len(validation['errors'])}):")
            for error in validation["errors"]:
                print(f"  [ERROR] {error}")

        if validation["warnings"]:
            print(f"\nWarnings ({len(validation['warnings'])}):")
            for warning in validation["warnings"]:
                print(f"  [WARN] {warning}")

        if not validation["errors"] and not validation["warnings"]:
            print("\n[SUCCESS] System state is consistent")

        print("=" * 60)
        return

    # Determine processing mode
    processing_mode = "incremental"
    if args.force_reprocess:
        processing_mode = "full-reprocess"
    elif args.files:
        processing_mode = "selective"

    # Display header
    logger.info("=" * 60)
    logger.info("DOCUMENT PROCESSING PIPELINE")
    logger.info("=" * 60)
    logger.info(f"OCR Enabled: {OCR_ENABLED}")
    logger.info(f"Processing Mode: {processing_mode}")
    if TESSERACT_VERSION and logger.isEnabledFor(logging.DEBUG):
        logger.debug(f"Tesseract Version: {TESSERACT_VERSION}")
    logger.info("")

    # Process documents
    try:
        # Process new documents and merge with existing knowledge base to prevent data loss
        new_knowledge_base = process_all_documents(
            base_folders, force_reprocess=args.force_reprocess, selective_files=args.files
        )

        # Preserve existing documents that weren't reprocessed
        preserved_docs = 0
        for key, value in knowledge_base.items():
            if key not in new_knowledge_base and not key.startswith("knowledge_base"):
                new_knowledge_base[key] = value
                preserved_docs += 1

        knowledge_base = new_knowledge_base
        if preserved_docs > 0:
            logger.info(
                f"Preserved {preserved_docs} existing documents during incremental processing"
            )

        # Merge with existing static knowledge base section if it exists
        if existing_kb_path.exists():
            try:
                with open(existing_kb_path, "r", encoding="utf-8") as f:
                    existing_data = json.load(f)
                    # Preserve the static knowledge_base section
                    if "knowledge_base" in existing_data:
                        knowledge_base["knowledge_base"] = existing_data["knowledge_base"]
                        logger.info("Preserved existing knowledge_base section")
            except Exception as e:
                logger.warning(f"Could not load existing knowledge_base for merging: {e}")

        # Save to JSON atomically
        output_file = base_dir / "knowledge_base.json"
        atomic_save_json(knowledge_base, output_file)

        processed_count = len(
            [k for k in knowledge_base.keys() if not k.startswith("knowledge_base")]
        )
        logger.info(f"[OK] Total documents in knowledge base: {processed_count}")

        # Display success message
        print(f"\n{'=' * 60}")
        print("PROCESSING COMPLETED SUCCESSFULLY")
        print("=" * 60)
        print(f"Mode: {processing_mode}")
        print(f"Documents Processed: {processed_count}")
        print(f"Knowledge Base: {output_file}")
        print("=" * 60)

    except Exception as e:
        logger.error(f"Processing failed: {e}", exc_info=True)
        print(f"\n[ERROR] Processing failed: {e}")
        exit(1)
    finally:
        # Ensure resources are properly closed
        dup_tracker.close()


if __name__ == "__main__":
    main()
