from docx import Document

def create_sample_rfp():
    doc = Document()
    doc.add_heading('REQUEST FOR PROPOSAL (RFP)', 0)
    
    doc.add_heading('1. Project Overview', level=1)
    doc.add_paragraph('Project Name: New Capital Water Treatment Plant Extension')
    doc.add_paragraph('Client: New Urban Communities Authority (NUCA)')
    doc.add_paragraph('Location: New Administrative Capital, Egypt')
    doc.add_paragraph('Duration: 18 months')
    
    doc.add_heading('2. Scope of Work', level=1)
    doc.add_paragraph('The consultant shall provide detailed engineering design for the extension of the existing water treatment plant.')
    doc.add_paragraph('The scope includes:')
    doc.add_paragraph('- Hydraulic design of sedimentation tanks.')
    doc.add_paragraph('- Structural design of reinforced concrete reservoirs.')
    doc.add_paragraph('- Electromechanical works for pumping stations.')
    doc.add_paragraph('- Geotechnical investigation and soil mechanics report.')
    
    doc.add_heading('3. Deliverables', level=1)
    doc.add_paragraph('- Inception Report')
    doc.add_paragraph('- Preliminary Design Report')
    doc.add_paragraph('- Final Design and Tender Documents')

    doc.save('sample_client_rfp.docx')
    print("Created sample_client_rfp.docx")

if __name__ == "__main__":
    create_sample_rfp()
