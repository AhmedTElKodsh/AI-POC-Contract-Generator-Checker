from docx import Document

def create_template():
    doc = Document()
    doc.add_heading('Engineering Proposal', 0)

    doc.add_paragraph('Project Name: {{ project_name }}')
    doc.add_paragraph('Client: {{ client }}')
    doc.add_paragraph('Location: {{ location }}')
    doc.add_paragraph('Date: {{ date }}')
    doc.add_paragraph('Duration: {{ duration_months }} months')

    doc.add_heading('1. Scope of Work', level=1)
    doc.add_paragraph('{{ sections.scope }}')

    doc.add_heading('2. Technical Glossary', level=1)
    p = doc.add_paragraph('The following terms are relevant to this project:')
    # In docxtpl, we can use loops
    doc.add_paragraph('{% for term, definition in metadata.glossary.items() %}')
    doc.add_paragraph('{{ term }}: {{ definition }}')
    doc.add_paragraph('{% endfor %}')

    doc.add_heading('3. Technical References', level=1)
    doc.add_paragraph('{% for ref in metadata.references %}')
    doc.add_paragraph('- {{ ref }}')
    doc.add_paragraph('{% endfor %}')

    doc.save('templates/poc_template.docx')
    print("Template created at templates/poc_template.docx")

if __name__ == "__main__":
    import os
    if not os.path.exists('templates'):
        os.makedirs('templates')
    create_template()
